import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import markdown
from datetime import datetime

def create_docx_summaries():
    """Create a nicely formatted DOCX file with all summaries"""
    doc = Document()
    
    # Add title
    title = doc.add_heading('Quicksilver - Book Summaries', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata
    doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d")}')
    doc.add_paragraph('Neal Stephenson - Quicksilver (The Baroque Cycle, Vol. 1)')
    doc.add_paragraph('')
    
    # Add table of contents
    doc.add_heading('Table of Contents', 1)
    doc.add_paragraph('Book One: Quicksilver', style='List Bullet')
    doc.add_paragraph('Book Two: King of the Vagabonds', style='List Bullet')
    doc.add_paragraph('Book Three: Odalisque', style='List Bullet')
    doc.add_page_break()
    
    # Process each book
    books = [
        ('Book One: Quicksilver', '/home/agentcode/text summaries/output/summaries/quicksilver_book_one_all_summaries.txt'),
        ('Book Two: King of the Vagabonds', '/home/agentcode/text summaries/output/summaries/quicksilver_book_two_all_summaries.txt'),
        ('Book Three: Odalisque', '/home/agentcode/text summaries/output/summaries/quicksilver_book_three_all_summaries.txt')
    ]
    
    for book_title, file_path in books:
        # Add book title
        doc.add_heading(book_title, 1)
        
        # Read summaries
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Parse summaries
        summaries = content.split('===')
        for summary in summaries[1:]:  # Skip first empty element
            if 'SUMMARY' in summary:
                lines = summary.strip().split('\n')
                # Extract summary number and word range
                header = lines[0].strip()
                word_count_line = lines[1] if len(lines) > 1 else ""
                summary_text = '\n'.join(lines[2:]).strip() if len(lines) > 2 else ""
                
                if summary_text:
                    # Add summary header
                    doc.add_heading(header.replace(' ===', ''), 2)
                    doc.add_paragraph(word_count_line, style='Intense Quote')
                    
                    # Add summary text
                    para = doc.add_paragraph(summary_text)
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    doc.add_paragraph('')  # Add spacing
        
        doc.add_page_break()
    
    # Save document
    doc.save('/home/agentcode/text summaries/output/Quicksilver_Book_Summaries.docx')
    print("Created: Quicksilver_Book_Summaries.docx")

def create_markdown_summaries():
    """Create a well-structured Markdown file with all summaries"""
    output = []
    
    # Add header
    output.append("# Quicksilver - Book Summaries\n")
    output.append(f"**Generated:** {datetime.now().strftime('%Y-%m-%d')}\n")
    output.append("**Book:** Neal Stephenson - Quicksilver (The Baroque Cycle, Vol. 1)\n")
    output.append("\n---\n")
    
    # Add table of contents
    output.append("## Table of Contents\n")
    output.append("- [Book One: Quicksilver](#book-one-quicksilver)")
    output.append("- [Book Two: King of the Vagabonds](#book-two-king-of-the-vagabonds)")
    output.append("- [Book Three: Odalisque](#book-three-odalisque)\n")
    output.append("\n---\n")
    
    # Process each book
    books = [
        ('Book One: Quicksilver', 'book-one-quicksilver', '/home/agentcode/text summaries/output/summaries/quicksilver_book_one_all_summaries.txt'),
        ('Book Two: King of the Vagabonds', 'book-two-king-of-the-vagabonds', '/home/agentcode/text summaries/output/summaries/quicksilver_book_two_all_summaries.txt'),
        ('Book Three: Odalisque', 'book-three-odalisque', '/home/agentcode/text summaries/output/summaries/quicksilver_book_three_all_summaries.txt')
    ]
    
    for book_title, book_id, file_path in books:
        # Add book section
        output.append(f"## {book_title}\n")
        
        # Read summaries
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Parse summaries
        summaries = content.split('===')
        for summary in summaries[1:]:  # Skip first empty element
            if 'SUMMARY' in summary:
                lines = summary.strip().split('\n')
                header = lines[0].strip().replace(' ===', '')
                word_count_line = lines[1] if len(lines) > 1 else ""
                summary_text = '\n'.join(lines[2:]).strip() if len(lines) > 2 else ""
                
                if summary_text:
                    # Add summary
                    output.append(f"### {header}\n")
                    output.append(f"*{word_count_line}*\n")
                    output.append(f"{summary_text}\n")
        
        output.append("\n---\n")
    
    # Save markdown file
    with open('/home/agentcode/text summaries/output/Quicksilver_Book_Summaries.md', 'w', encoding='utf-8') as f:
        f.write('\n'.join(output))
    
    print("Created: Quicksilver_Book_Summaries.md")

def create_html_summaries():
    """Create an HTML file with navigation and styling"""
    html_template = """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Quicksilver - Book Summaries</title>
    <style>
        body {
            font-family: Georgia, serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            background-color: #f5f5f5;
        }
        .container {
            background-color: white;
            padding: 40px;
            box-shadow: 0 0 10px rgba(0,0,0,0.1);
        }
        h1 {
            color: #2c3e50;
            text-align: center;
            border-bottom: 3px solid #3498db;
            padding-bottom: 10px;
        }
        h2 {
            color: #34495e;
            margin-top: 40px;
            border-bottom: 2px solid #ecf0f1;
            padding-bottom: 5px;
        }
        h3 {
            color: #7f8c8d;
            margin-top: 30px;
        }
        .metadata {
            text-align: center;
            color: #7f8c8d;
            margin-bottom: 30px;
        }
        .toc {
            background-color: #ecf0f1;
            padding: 20px;
            border-radius: 5px;
            margin-bottom: 30px;
        }
        .toc ul {
            list-style-type: none;
            padding-left: 20px;
        }
        .toc a {
            color: #3498db;
            text-decoration: none;
        }
        .toc a:hover {
            text-decoration: underline;
        }
        .summary {
            text-align: justify;
            margin-bottom: 20px;
        }
        .word-count {
            font-style: italic;
            color: #7f8c8d;
            margin-bottom: 10px;
        }
        .back-to-top {
            position: fixed;
            bottom: 20px;
            right: 20px;
            background-color: #3498db;
            color: white;
            padding: 10px 15px;
            border-radius: 5px;
            text-decoration: none;
            display: none;
        }
    </style>
</head>
<body>
    <div class="container">
        <h1>Quicksilver - Book Summaries</h1>
        <div class="metadata">
            <p><strong>Generated:</strong> {date}</p>
            <p><strong>Book:</strong> Neal Stephenson - Quicksilver (The Baroque Cycle, Vol. 1)</p>
        </div>
        
        <div class="toc">
            <h2>Table of Contents</h2>
            <ul>
                <li><a href="#book-one">Book One: Quicksilver</a></li>
                <li><a href="#book-two">Book Two: King of the Vagabonds</a></li>
                <li><a href="#book-three">Book Three: Odalisque</a></li>
            </ul>
        </div>
        
        {content}
        
        <a href="#top" class="back-to-top">↑ Top</a>
    </div>
    
    <script>
        // Show/hide back to top button
        window.onscroll = function() {
            if (document.body.scrollTop > 200 || document.documentElement.scrollTop > 200) {
                document.querySelector('.back-to-top').style.display = "block";
            } else {
                document.querySelector('.back-to-top').style.display = "none";
            }
        };
    </script>
</body>
</html>"""
    
    content_parts = []
    
    # Process each book
    books = [
        ('Book One: Quicksilver', 'book-one', '/home/agentcode/text summaries/output/summaries/quicksilver_book_one_all_summaries.txt'),
        ('Book Two: King of the Vagabonds', 'book-two', '/home/agentcode/text summaries/output/summaries/quicksilver_book_two_all_summaries.txt'),
        ('Book Three: Odalisque', 'book-three', '/home/agentcode/text summaries/output/summaries/quicksilver_book_three_all_summaries.txt')
    ]
    
    for book_title, book_id, file_path in books:
        content_parts.append(f'<h2 id="{book_id}">{book_title}</h2>')
        
        # Read summaries
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Parse summaries
        summaries = content.split('===')
        for summary in summaries[1:]:
            if 'SUMMARY' in summary:
                lines = summary.strip().split('\n')
                header = lines[0].strip().replace(' ===', '')
                word_count_line = lines[1] if len(lines) > 1 else ""
                summary_text = '\n'.join(lines[2:]).strip() if len(lines) > 2 else ""
                
                if summary_text:
                    content_parts.append(f'<h3>{header}</h3>')
                    content_parts.append(f'<p class="word-count">{word_count_line}</p>')
                    content_parts.append(f'<p class="summary">{summary_text}</p>')
    
    # Generate final HTML
    html_content = html_template.format(
        date=datetime.now().strftime('%Y-%m-%d'),
        content='\n'.join(content_parts)
    )
    
    # Save HTML file
    with open('/home/agentcode/text summaries/output/Quicksilver_Book_Summaries.html', 'w', encoding='utf-8') as f:
        f.write(html_content)
    
    print("Created: Quicksilver_Book_Summaries.html")

if __name__ == "__main__":
    print("Creating readable formats...")
    
    # Check if python-docx is available
    try:
        create_docx_summaries()
    except ImportError:
        print("Note: python-docx not installed. Skipping DOCX creation.")
        print("To install: pip install python-docx")
    
    # Create Markdown (always works)
    create_markdown_summaries()
    
    # Create HTML (always works)
    create_html_summaries()
    
    print("\nFormats created successfully!")
    print("- Markdown (.md): Great for GitHub, documentation, and text editors")
    print("- HTML: Perfect for web viewing with navigation")
    print("- DOCX: Best for Word users and formal documents")